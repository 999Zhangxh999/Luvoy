# ============================================
# 简历解析服务（Agent 2）
# 负责：简历上传解析、学生就业能力画像生成
# 技术亮点：多格式支持 / pdfplumber文本提取 / OCR扫描件识别
# ============================================
import json
import os
import logging

from models.database import db
from models.student import Student, StudentProfile
from services.llm_client import get_llm
from prompts.templates import RESUME_PARSE_SYSTEM, RESUME_PARSE_USER

logger = logging.getLogger(__name__)

# OCR支持检测与配置
OCR_AVAILABLE = False
USE_PYMUPDF = False  # 使用PyMuPDF渲染PDF（无需Poppler）

# 检测PyMuPDF（推荐：纯Python，无外部依赖）
try:
    import fitz  # PyMuPDF
    USE_PYMUPDF = True
    logger.info("PyMuPDF可用，将用于PDF图像渲染")
except ImportError:
    logger.info("PyMuPDF未安装，将尝试pdf2image+Poppler")

# 检测Tesseract OCR
try:
    import pytesseract
    from PIL import Image
    
    # 配置Tesseract路径（Windows常见安装位置）
    tesseract_paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        os.path.expanduser(r"~\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"),
    ]
    for tess_path in tesseract_paths:
        if os.path.exists(tess_path):
            pytesseract.pytesseract.tesseract_cmd = tess_path
            logger.info(f"Tesseract配置成功: {tess_path}")
            break
    
    # 验证Tesseract是否可用
    try:
        pytesseract.get_tesseract_version()
        if USE_PYMUPDF:
            OCR_AVAILABLE = True
            logger.info("OCR功能就绪（PyMuPDF + Tesseract）")
        else:
            # 备选：尝试pdf2image+Poppler
            try:
                from pdf2image import convert_from_path
                import shutil
                if shutil.which('pdftoppm'):
                    OCR_AVAILABLE = True
                    logger.info("OCR功能就绪（pdf2image + Poppler + Tesseract）")
            except ImportError:
                pass
    except Exception as e:
        logger.warning(f"Tesseract不可用: {e}")
        
except ImportError as e:
    logger.info(f"OCR依赖未安装(pytesseract)，扫描版PDF将无法解析: {e}")


class ResumeParser:
    """简历解析Agent - 从简历中提取学生就业能力画像"""

    def _extract_skills_from_text(self, text: str) -> list:
        """从文本中提取基础技能关键词，作为LLM不可用时的降级方案。"""
        if not text:
            return []

        keywords = [
            'Python', 'Java', 'JavaScript', 'TypeScript', 'C++', 'SQL',
            'Vue', 'React', 'Flask', 'Django', 'Spring', 'Node.js',
            '机器学习', '深度学习', '数据分析', '算法', 'Linux', 'Docker',
            'Git', 'Redis', 'MySQL', 'TensorFlow', 'PyTorch'
        ]

        lowered = text.lower()
        found = []
        for kw in keywords:
            if kw.lower() in lowered:
                found.append(kw)
        return found[:12]

    def _build_fallback_profile_result(self, student: Student) -> dict:
        """构建离线降级画像，保证流程可继续。"""
        resume_text = student.resume_text or ''

        technical_skills = self._extract_skills_from_text(resume_text)
        if student.major and '计算机' in student.major and 'Python' not in technical_skills:
            technical_skills.append('Python')
        if student.major and ('人工智能' in student.major or '智能' in student.major) and '机器学习' not in technical_skills:
            technical_skills.append('机器学习')

        completeness_base = 45
        if student.name:
            completeness_base += 8
        if student.education:
            completeness_base += 8
        if student.major:
            completeness_base += 8
        if student.school:
            completeness_base += 8
        if resume_text:
            completeness_base += min(20, max(5, len(resume_text) // 300))

        competitiveness = min(88, 40 + len(technical_skills) * 4)

        strengths = []
        if technical_skills:
            strengths.append('具备一定技术基础')
        if resume_text and len(resume_text) > 500:
            strengths.append('项目/经历描述较完整')
        if student.education in ('本科', '硕士', '博士'):
            strengths.append('学历背景具备竞争力')
        if not strengths:
            strengths = ['基础信息完整，可继续补充经历提升画像准确度']

        weaknesses = [
            '建议补充可量化的项目成果',
            '建议明确目标岗位并强化对应技能'
        ]

        return {
            'technical_skills': technical_skills,
            'certificates': [],
            'project_experience': [],
            'internship_experience': [],
            'awards': [],
            'innovation_ability': 6,
            'learning_ability': 7,
            'pressure_resistance': 6,
            'communication_skill': 6,
            'teamwork_ability': 6,
            'internship_ability': 5,
            'completeness_score': min(100, completeness_base),
            'competitiveness_score': competitiveness,
            'overall_evaluation': '当前为离线降级画像（LLM服务暂不可用），可用于继续匹配流程，建议稍后重试以获得更精准分析。',
            'strengths': strengths,
            'weaknesses': weaknesses,
            'basic_info': {
                'name': student.name,
                'education': student.education,
                'major': student.major,
                'school': student.school,
            },
        }
    
    def parse_pdf(self, file_path: str) -> str:
        """
        解析PDF简历 - 多策略解析
        
        策略优先级：
        1. pdfplumber（文本PDF首选）
        2. PyPDF2（备选）
        3. OCR（扫描版PDF）
        """
        text = ''
        
        # 策略1：使用pdfplumber（对文本PDF效果最好）
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + '\n'
            if text.strip():
                logger.info(f"PDF解析成功(pdfplumber): {len(text)} 字符")
                return text.strip()
        except Exception as e:
            logger.warning(f"pdfplumber解析失败: {e}")
        
        # 策略2：PyPDF2备选
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + '\n'
            if text.strip():
                logger.info(f"PDF解析成功(PyPDF2): {len(text)} 字符")
                return text.strip()
        except Exception as e:
            logger.warning(f"PyPDF2解析失败: {e}")
        
        # 策略3：OCR（扫描版PDF）
        if OCR_AVAILABLE:
            try:
                text = self._ocr_pdf(file_path)
                if text.strip():
                    logger.info(f"PDF解析成功(OCR): {len(text)} 字符")
                    return text.strip()
            except Exception as e:
                logger.warning(f"OCR解析失败: {e}")
        
        # 所有策略都失败
        logger.error(f"PDF解析失败: 可能是扫描版PDF且OCR不可用")
        return ''
    
    def _ocr_pdf(self, file_path: str) -> str:
        """使用OCR识别扫描版PDF"""
        if not OCR_AVAILABLE:
            return ''
        
        text = ''
        
        # 方法1：使用PyMuPDF渲染PDF为图像（推荐，无需Poppler）
        if USE_PYMUPDF:
            try:
                import fitz
                doc = fitz.open(file_path)
                for i, page in enumerate(doc):
                    # 渲染页面为图像（200 DPI）
                    mat = fitz.Matrix(200/72, 200/72)
                    pix = page.get_pixmap(matrix=mat)
                    
                    # 转换为PIL Image
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    
                    # OCR识别
                    page_text = pytesseract.image_to_string(img, lang='chi_sim+eng')
                    text += page_text + '\n'
                    logger.info(f"OCR处理第{i+1}页完成（PyMuPDF）")
                doc.close()
                return text.strip()
            except Exception as e:
                logger.warning(f"PyMuPDF OCR失败: {e}")
        
        # 方法2：使用pdf2image + Poppler（备选）
        try:
            from pdf2image import convert_from_path
            images = convert_from_path(file_path, dpi=200)
            for i, image in enumerate(images):
                page_text = pytesseract.image_to_string(image, lang='chi_sim+eng')
                text += page_text + '\n'
                logger.info(f"OCR处理第{i+1}页完成（pdf2image）")
            return text.strip()
        except Exception as e:
            logger.warning(f"pdf2image OCR失败: {e}")
        
        return text.strip()
    
    def parse_docx(self, file_path: str) -> str:
        """解析Word简历"""
        try:
            from docx import Document
            doc = Document(file_path)
            
            # 提取段落文本
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # 提取表格中的文本（简历常用表格格式）
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        paragraphs.append(row_text)
            
            text = '\n'.join(paragraphs)
            logger.info(f"DOCX解析成功: {len(text)} 字符")
            return text.strip()
        except Exception as e:
            logger.error(f"DOCX解析失败: {e}")
            return ''
    
    def parse_file(self, file_path: str) -> str:
        """根据文件类型自动选择解析器"""
        ext = os.path.splitext(file_path)[1].lower()
        logger.info(f"解析文件: {file_path}, 扩展名: {ext}")
        
        if ext == '.pdf':
            return self.parse_pdf(file_path)
        elif ext == '.docx':
            return self.parse_docx(file_path)
        elif ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read().strip()
        elif not ext:
            # 无扩展名，尝试作为PDF解析
            logger.warning(f"文件无扩展名，尝试按PDF解析: {file_path}")
            return self.parse_pdf(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")
    
    def generate_student_profile(self, student: Student) -> StudentProfile:
        """
        根据学生信息和简历内容，生成就业能力画像
        
        Args:
            student: 学生基本信息（包含简历文本）
        Returns:
            StudentProfile: 生成的能力画像
        """
        llm = get_llm()
        
        # 构建学生基础信息
        student_info = f"""
姓名：{student.name or '未知'}
性别：{student.gender or '未知'}
年龄：{student.age or '未知'}
学历：{student.education or '未知'}
专业：{student.major or '未知'}
学校：{student.school or '未知'}
期望岗位：{', '.join(student.get_target_positions()) if student.get_target_positions() else '未设定'}
期望薪资：{student.salary_expectation or '未设定'}
"""
        
        resume_text = student.resume_text or '暂无简历内容，请根据基础信息进行评估。'
        
        messages = [
            {"role": "system", "content": RESUME_PARSE_SYSTEM},
            {"role": "user", "content": RESUME_PARSE_USER.format(
                student_info=student_info,
                resume_text=resume_text,
            )}
        ]
        
        try:
            result = llm.chat_json(messages, temperature=0.3)
        except Exception as e:
            logger.warning(f"LLM画像生成失败，启用降级画像: {e}")
            result = self._build_fallback_profile_result(student)
        
        # 创建或更新学生画像
        profile = StudentProfile.query.filter_by(student_id=student.id).first()
        if not profile:
            profile = StudentProfile(student_id=student.id)
        
        # 填充技术技能
        tech_skills = result.get('technical_skills', [])
        profile.technical_skills = json.dumps(tech_skills, ensure_ascii=False)
        
        # 填充证书
        profile.certificates = json.dumps(
            result.get('certificates', []), ensure_ascii=False
        )
        
        # 填充项目经历
        profile.project_experience = json.dumps(
            result.get('project_experience', []), ensure_ascii=False
        )
        
        # 填充实习经历
        profile.internship_experience = json.dumps(
            result.get('internship_experience', []), ensure_ascii=False
        )
        
        # 填充获奖
        profile.awards = json.dumps(
            result.get('awards', []), ensure_ascii=False
        )
        
        # 软技能评分
        profile.innovation_ability = result.get('innovation_ability', 5)
        profile.learning_ability = result.get('learning_ability', 5)
        profile.pressure_resistance = result.get('pressure_resistance', 5)
        profile.communication_skill = result.get('communication_skill', 5)
        profile.teamwork_ability = result.get('teamwork_ability', 5)
        profile.internship_ability = result.get('internship_ability', 5)
        
        # 综合评分
        profile.completeness_score = result.get('completeness_score', 50)
        profile.competitiveness_score = result.get('competitiveness_score', 50)
        
        # 评价
        profile.overall_evaluation = result.get('overall_evaluation', '')
        profile.strengths = json.dumps(
            result.get('strengths', []), ensure_ascii=False
        )
        profile.weaknesses = json.dumps(
            result.get('weaknesses', []), ensure_ascii=False
        )
        
        # 更新学生基础信息（如果LLM提取到了更多信息）
        basic_info = result.get('basic_info', {})
        if basic_info.get('name') and not student.name:
            student.name = basic_info['name']
        if basic_info.get('education') and not student.education:
            student.education = basic_info['education']
        if basic_info.get('major') and not student.major:
            student.major = basic_info['major']
        if basic_info.get('school') and not student.school:
            student.school = basic_info['school']
        
        db.session.add(profile)
        db.session.commit()
        
        logger.info(f"学生画像生成成功: {student.name}, 完整度={profile.completeness_score}, 竞争力={profile.competitiveness_score}")
        return profile
    
    def create_student_from_form(self, form_data: dict) -> Student:
        """
        从表单数据创建学生记录
        
        Args:
            form_data: 表单提交的数据
        Returns:
            Student: 创建的学生记录
        """
        student = Student(
            name=form_data.get('name', ''),
            gender=form_data.get('gender', ''),
            age=int(form_data['age']) if form_data.get('age') else None,
            education=form_data.get('education', ''),
            major=form_data.get('major', ''),
            school=form_data.get('school', ''),
            graduation_year=int(form_data['graduation_year']) if form_data.get('graduation_year') else None,
            email=form_data.get('email', ''),
            phone=form_data.get('phone', ''),
            resume_text=form_data.get('resume_text', ''),
            target_positions=json.dumps(
                [p.strip() for p in form_data.get('target_positions', '').split(',') if p.strip()],
                ensure_ascii=False
            ),
            target_cities=json.dumps(
                [c.strip() for c in form_data.get('target_cities', '').split(',') if c.strip()],
                ensure_ascii=False
            ),
            salary_expectation=form_data.get('salary_expectation', ''),
            career_preference=form_data.get('career_preference', ''),
        )
        
        db.session.add(student)
        db.session.commit()
        
        logger.info(f"学生记录创建成功: {student.name} (ID={student.id})")
        return student
